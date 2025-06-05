import os
from flask import Flask, request, render_template, redirect, url_for, send_from_directory, flash, session, send_file
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import io # Required for handling image bytes, BytesIO
import uuid # For generating unique filenames
import zipfile # For creating ZIP archives
from docx import Document # For creating .docx files
from docx.shared import Inches, Pt # For image sizing and font
from docx.enum.text import WD_ALIGN_PARAGRAPH # For text alignment

# Configuração inicial
UPLOAD_FOLDER = 'uploads'
EXTRACTED_IMAGES_FOLDER = 'extracted_images'
ALLOWED_EXTENSIONS = {'pdf'}

# CONFIGURAR CAMINHO DO TESSERACT SE NECESSÁRIO (descomente e ajuste a linha abaixo)
# No Windows, por exemplo: r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# No Linux, geralmente é detectado automaticamente ou pode ser /usr/bin/tesseract
# pytesseract.pytesseract.tesseract_cmd = r'<caminho_para_seu_tesseract_exe>'

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['EXTRACTED_IMAGES_FOLDER'] = EXTRACTED_IMAGES_FOLDER
app.secret_key = "super secret key" # Necessário para flash messages, etc.

# Certificar que os diretórios de upload e imagens existem
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(EXTRACTED_IMAGES_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'pdf_file' not in request.files:
            flash('Nenhum arquivo selecionado.', 'error')
            return redirect(request.url)
        file = request.files['pdf_file']
        if file.filename == '':
            flash('Nenhum arquivo selecionado.', 'error')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            original_filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], original_filename)
            file.save(filepath)
            
            flash(f'Arquivo "{original_filename}" recebido. Processando...', 'info')
            try:
                extracted_text, image_files = process_pdf(filepath, original_filename)
                flash('Processamento concluído!', 'success')
                
                session[f'{original_filename}_images'] = image_files
                session[f'{original_filename}_pdf_path'] = filepath

                return render_template('index.html', 
                                       filename=original_filename, 
                                       extracted_text=extracted_text, 
                                       images=image_files,
                                       processed=True)
            except Exception as e:
                flash(f'Erro ao processar o PDF: {str(e)}', 'error')
                if os.path.exists(filepath):
                    os.remove(filepath)
                return redirect(request.url)
        else:
            flash('Formato de arquivo inválido. Por favor, envie um .pdf.', 'error')
            return redirect(request.url)
            
    return render_template('index.html', processed=False)

def process_pdf(pdf_filepath, original_filename):
    doc = fitz.open(pdf_filepath)
    full_text = []
    extracted_image_files = []
    image_counter = 1 

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)
        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            unique_img_filename = f"{original_filename}_p{page_num+1}_img{image_counter}.{image_ext}"
            image_save_path = os.path.join(app.config['EXTRACTED_IMAGES_FOLDER'], unique_img_filename)
            
            with open(image_save_path, "wb") as img_file:
                img_file.write(image_bytes)
            extracted_image_files.append({'filename': unique_img_filename, 'path': image_save_path})
            image_counter += 1

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        page_text = page.get_text("text")
        
        if not page_text.strip(): 
            try:
                images_from_page = convert_from_path(pdf_filepath, dpi=300, first_page=page_num + 1, last_page=page_num + 1)
                if images_from_page:
                    ocr_text = pytesseract.image_to_string(images_from_page[0], lang='por') 
                    full_text.append(f"--- OCR Página {page_num + 1} ---\n{ocr_text}")
                else:
                    full_text.append(f"--- Página {page_num + 1}: Não foi possível converter para imagem para OCR ---")
            except Exception as e:
                print(f"Erro no OCR da página {page_num + 1}: {e}")
                full_text.append(f"--- Página {page_num + 1}: Erro durante OCR ({e}) ---")
        else:
            full_text.append(page_text)
            
    doc.close()
    return "\n\n".join(full_text), extracted_image_files

@app.route('/extracted_images/<filename>')
def send_extracted_image(filename):
    return send_from_directory(app.config['EXTRACTED_IMAGES_FOLDER'], filename)

@app.route('/download/<original_pdf_name>', methods=['POST'])
def download_result(original_pdf_name):
    download_type = request.form.get('download_type')
    text_to_download = request.form.get('extracted_text_for_download', '')
    
    image_files_info = session.get(f'{original_pdf_name}_images', [])
    filename_base = original_pdf_name.rsplit('.', 1)[0]

    if download_type == 'docx':
        try:
            doc = Document()
            paragraphs = text_to_download.split('\n\n') 
            for para_text in paragraphs:
                if para_text.strip(): 
                    p = doc.add_paragraph()
                    lines = para_text.split('\n')
                    for i, line_text in enumerate(lines):
                        p.add_run(line_text)
                        if i < len(lines) - 1:
                            p.add_run().add_break()
            
            if image_files_info:
                doc.add_heading('Imagens Extraídas', level=2)
                for img_info in image_files_info:
                    try:
                        doc.add_picture(img_info['path'], width=Inches(4.0)) 
                    except Exception as e:
                        print(f"Erro ao adicionar imagem {img_info['filename']} ao DOCX: {e}")
                        doc.add_paragraph(f"[Erro ao carregar imagem: {img_info['filename']}]", style='Emphasis')

            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            return send_file(file_stream, as_attachment=True, download_name=f'{filename_base}_extraido.docx',
                             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        except Exception as e:
            flash(f'Erro ao gerar arquivo .docx: {e}', 'error')
            return redirect(url_for('index'))

    elif download_type == 'txt_images':
        try:
            zip_stream = io.BytesIO()
            with zipfile.ZipFile(zip_stream, 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(f'{filename_base}_texto.txt', text_to_download)
                if image_files_info:
                    for img_info in image_files_info:
                        try:
                            zf.write(img_info['path'], arcname=f"imagens/{img_info['filename']}")
                        except Exception as e:
                            print(f"Erro ao adicionar imagem {img_info['filename']} ao ZIP: {e}")
            
            zip_stream.seek(0)
            return send_file(zip_stream, as_attachment=True, download_name=f'{filename_base}_extraido.zip',
                             mimetype='application/zip')
        except Exception as e:
            flash(f'Erro ao gerar arquivo .zip: {e}', 'error')
            return redirect(url_for('index'))
    else:
        flash('Tipo de download inválido.', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
